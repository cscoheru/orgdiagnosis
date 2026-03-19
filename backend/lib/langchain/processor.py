"""
Document Processor for Five-Dimensional Diagnosis

This module handles document loading, cleaning, and semantic chunking
using LangChain's document loaders and text splitters.
"""

import os
import re
from typing import List, Optional, Tuple
from loguru import logger

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
)


class DocumentProcessor:
    """
    文档处理器：负责文档加载、清洗、分片

    核心功能：
    1. 支持 PDF/Word/TXT 文件
    2. 按语义切分（优先按 Markdown 标题）
    3. 保留文档层级结构
    4. 关键词预过滤
    """

    # 维度关键词映射（用于预过滤）
    DIMENSION_KEYWORDS = {
        "strategy": [
            "战略", "规划", "目标", "愿景", "使命", "市场", "竞争", "业务", "增长",
            "业绩", "机会", "创新", "执行", "关键任务", "战略意图", "经营分析"
        ],
        "structure": [
            "组织", "架构", "部门", "层级", "边界", "决策", "授权", "岗位",
            "流程", "协作", "数字化", "人效", "响应", "效率", "汇报"
        ],
        "performance": [
            "绩效", "考核", "目标", "KPI", "OKR", "指标", "辅导", "反馈",
            "公平", "申诉", "激励", "晋升", "淘汰", "培训", "发展"
        ],
        "compensation": [
            "薪酬", "工资", "奖金", "激励", "福利", "调薪", "定位",
            "固浮比", "公平", "总额", "沟通", "短期", "长期"
        ],
        "talent": [
            "人才", "招聘", "盘点", "梯队", "胜任力", "雇主", "品牌",
            "培养", "骨干", "职业", "流失", "敬业", "非物质", "流动"
        ]
    }

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        enable_keyword_filter: bool = True
    ):
        """
        初始化文档处理器

        Args:
            chunk_size: 分片大小（字符数）
            chunk_overlap: 分片重叠（字符数）
            enable_keyword_filter: 是否启用关键词预过滤
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_keyword_filter = enable_keyword_filter

        # 初始化分片器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
            length_function=len,
        )

        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "h1"),
                ("##", "h2"),
                ("###", "h3"),
            ]
        )

        logger.info(f"DocumentProcessor initialized: chunk_size={chunk_size}, overlap={chunk_overlap}")

    def load_text(self, text: str, source: str = "input") -> List[Document]:
        """
        从纯文本创建文档

        Args:
            text: 原始文本
            source: 来源标识

        Returns:
            Document 列表
        """
        doc = Document(page_content=text, metadata={"source": source})
        return [doc]

    def load_from_file(self, file_path: str) -> List[Document]:
        """
        从文件加载文档

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext in [".doc", ".docx"]:
            return self._load_word(file_path)
        elif ext == ".txt":
            return self._load_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文件"""
        try:
            # 使用 PyMuPDF (fitz)
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            documents = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": file_path,
                            "page": page_num + 1,
                            "type": "pdf"
                        }
                    ))

            doc.close()
            logger.info(f"Loaded {len(documents)} pages from PDF: {file_path}")
            return documents

        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to basic text extraction")
            # Fallback: 使用 pdfplumber
            import pdfplumber
            documents = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": page_num + 1,
                                "type": "pdf"
                            }
                        ))

            logger.info(f"Loaded {len(documents)} pages from PDF (pdfplumber): {file_path}")
            return documents

    def _load_word(self, file_path: str) -> List[Document]:
        """加载 Word 文件"""
        try:
            # 尝试使用 unstructured
            from unstructured.partition.docx import partition_docx
            elements = partition_docx(filename=file_path)

            documents = []
            for element in elements:
                text = str(element)
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": file_path,
                            "type": "docx"
                        }
                    ))

            logger.info(f"Loaded {len(documents)} sections from Word: {file_path}")
            return documents

        except ImportError:
            logger.warning("unstructured not installed, falling back to python-docx")
            # Fallback: 使用 python-docx
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)

            documents = []
            for para in doc.paragraphs:
                if para.text.strip():
                    documents.append(Document(
                        page_content=para.text,
                        metadata={
                            "source": file_path,
                            "type": "docx"
                        }
                    ))

            logger.info(f"Loaded {len(documents)} paragraphs from Word: {file_path}")
            return documents

    def _load_text_file(self, file_path: str) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        return [Document(
            page_content=text,
            metadata={
                "source": file_path,
                "type": "txt"
            }
        )]

    def clean_text(self, documents: List[Document]) -> List[Document]:
        """
        清洗文本

        - 去除多余空白
        - 去除页眉页脚模式
        - 保留有意义的换行

        Args:
            documents: 原始文档列表

        Returns:
            清洗后的文档列表
        """
        cleaned = []

        for doc in documents:
            text = doc.page_content

            # 去除页码模式
            text = re.sub(r'\n\s*第?\s*\d+\s*页?\s*\n', '\n', text)

            # 去除重复的空行
            text = re.sub(r'\n{3,}', '\n\n', text)

            # 去除行首行尾空白
            lines = [line.strip() for line in text.split('\n')]
            text = '\n'.join(lines)

            # 去除多余的空格（保留中文之间的空格用于断句）
            text = re.sub(r'[ \t]{2,}', ' ', text)

            if text.strip():
                cleaned.append(Document(
                    page_content=text,
                    metadata=doc.metadata
                ))

        logger.info(f"Cleaned {len(documents)} documents -> {len(cleaned)} documents")
        return cleaned

    def split_documents(
        self,
        documents: List[Document],
        preserve_structure: bool = True
    ) -> List[Document]:
        """
        分片文档

        Args:
            documents: 文档列表
            preserve_structure: 是否保留层级结构（Markdown 标题）

        Returns:
            分片后的文档列表
        """
        all_chunks = []

        for doc in documents:
            text = doc.page_content

            # 检测是否包含 Markdown 标题
            has_markdown_headers = bool(re.search(r'^#{1,3}\s+', text, re.MULTILINE))

            if preserve_structure and has_markdown_headers:
                # 按 Markdown 标题切分
                try:
                    md_chunks = self.markdown_splitter.split_text(text)
                    for chunk in md_chunks:
                        all_chunks.append(Document(
                            page_content=chunk.page_content,
                            metadata={**doc.metadata, **chunk.metadata}
                        ))
                except Exception as e:
                    logger.warning(f"Markdown splitting failed: {e}, falling back to recursive")
                    chunks = self.text_splitter.split_text(text)
                    for chunk in chunks:
                        all_chunks.append(Document(
                            page_content=chunk,
                            metadata=doc.metadata
                        ))
            else:
                # 递归字符切分
                chunks = self.text_splitter.split_text(text)
                for chunk in chunks:
                    all_chunks.append(Document(
                        page_content=chunk,
                        metadata=doc.metadata
                    ))

        logger.info(f"Split {len(documents)} documents -> {len(all_chunks)} chunks")
        return all_chunks

    def filter_by_dimension(
        self,
        documents: List[Document],
        dimension: str,
        min_matches: int = 1
    ) -> List[Document]:
        """
        按维度关键词过滤文档

        Args:
            documents: 文档列表
            dimension: 维度名称 (strategy/structure/performance/compensation/talent)
            min_matches: 最少匹配关键词数

        Returns:
            过滤后的文档列表
        """
        if not self.enable_keyword_filter:
            return documents

        if dimension not in self.DIMENSION_KEYWORDS:
            logger.warning(f"Unknown dimension: {dimension}, returning all documents")
            return documents

        keywords = self.DIMENSION_KEYWORDS[dimension]
        filtered = []

        for doc in documents:
            text = doc.page_content
            matches = sum(1 for kw in keywords if kw in text)

            if matches >= min_matches:
                filtered.append(doc)

        logger.info(f"Filtered by '{dimension}': {len(documents)} -> {len(filtered)} documents")
        return filtered

    def get_relevant_chunks(
        self,
        documents: List[Document],
        dimension: str,
        max_chunks: int = 10
    ) -> List[Tuple[Document, int]]:
        """
        获取与维度最相关的文档片段

        按关键词匹配数排序，返回最相关的片段

        Args:
            documents: 文档列表
            dimension: 维度名称
            max_chunks: 最大返回数量

        Returns:
            (Document, score) 列表，按分数降序
        """
        if dimension not in self.DIMENSION_KEYWORDS:
            return [(doc, 0) for doc in documents[:max_chunks]]

        keywords = self.DIMENSION_KEYWORDS[dimension]
        scored = []

        for doc in documents:
            text = doc.page_content
            score = sum(1 for kw in keywords if kw in text)
            scored.append((doc, score))

        # 按分数降序排序
        scored.sort(key=lambda x: x[1], reverse=True)

        return scored[:max_chunks]

    def process(
        self,
        text: Optional[str] = None,
        file_path: Optional[str] = None,
        clean: bool = True,
        split: bool = True,
        chunk_size: Optional[int] = None
    ) -> List[Document]:
        """
        完整的文档处理流程

        Args:
            text: 原始文本（二选一）
            file_path: 文件路径（二选一）
            clean: 是否清洗
            split: 是否分片
            chunk_size: 临时指定分片大小

        Returns:
            处理后的文档列表
        """
        # 加载
        if text:
            documents = self.load_text(text)
        elif file_path:
            documents = self.load_from_file(file_path)
        else:
            raise ValueError("Either text or file_path must be provided")

        # 清洗
        if clean:
            documents = self.clean_text(documents)

        # 分片
        if split:
            old_chunk_size = self.text_splitter._chunk_size
            if chunk_size:
                self.text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
                )
            documents = self.split_documents(documents)
            if chunk_size:
                self.text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=old_chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
                )

        return documents


# 便捷函数
def process_text(
    text: str,
    chunk_size: int = 1000,
    dimension: Optional[str] = None
) -> List[Document]:
    """
    处理文本的便捷函数

    Args:
        text: 原始文本
        chunk_size: 分片大小
        dimension: 可选的维度过滤

    Returns:
        处理后的文档列表
    """
    processor = DocumentProcessor(chunk_size=chunk_size)
    documents = processor.process(text=text)

    if dimension:
        documents = processor.filter_by_dimension(documents, dimension)

    return documents
